"""
Document ingestion pipeline for Akasha RAG system.

This module handles loading, parsing, and chunking of various document formats
for embedding and storage in the vector database.
"""

import hashlib
import mimetypes
import tempfile
import json
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Iterator, Tuple
import asyncio
import time
import base64
from io import BytesIO

from pydantic import BaseModel, Field
import numpy as np

from ..core.logging import get_logger, PerformanceLogger
from ..core.exceptions import AkashaError


class DocumentFormat(str, Enum):
    """Supported document formats."""
    TEXT = "text"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"
    CSV = "csv"
    IMAGE = "image"
    OTHER = "other"


class ChunkingStrategy(str, Enum):
    """Document chunking strategies."""
    FIXED_SIZE = "fixed_size"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"
    SEMANTIC = "semantic"
    RECURSIVE = "recursive"
    LAYOUT_AWARE = "layout_aware"


@dataclass
class DocumentMetadata:
    """Metadata for processed documents."""
    document_id: str
    file_path: str
    file_name: str
    file_size: int
    file_hash: str
    mime_type: str
    format: DocumentFormat
    processed_at: float
    chunk_count: int
    processing_time: float
    source: str = "local"
    custom_metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.custom_metadata is None:
            self.custom_metadata = {}


class ContentType(str, Enum):
    """Types of content within documents."""
    TEXT = "text"
    TITLE = "title"
    HEADER = "header"
    PARAGRAPH = "paragraph"
    LIST = "list"
    TABLE = "table"
    IMAGE = "image"
    FORMULA = "formula"
    FOOTNOTE = "footnote"
    CAPTION = "caption"


class DocumentChunk(BaseModel):
    """A chunk of document content with metadata."""
    id: str = Field(..., description="Unique chunk identifier")
    content: str = Field(..., description="Chunk text content")
    document_id: str = Field(..., description="Parent document identifier")
    chunk_index: int = Field(..., description="Index within document")
    start_offset: int = Field(default=0, description="Start character offset")
    end_offset: int = Field(default=0, description="End character offset")
    content_type: ContentType = Field(default=ContentType.TEXT, description="Type of content")
    page_number: Optional[int] = Field(default=None, description="Source page number")
    bbox: Optional[List[float]] = Field(default=None, description="Bounding box coordinates")
    confidence: float = Field(default=1.0, description="Extraction confidence score")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Chunk metadata")
    embedding: Optional[List[float]] = Field(default=None, description="Chunk embedding vector")
    image_data: Optional[str] = Field(default=None, description="Base64 encoded image data")
    
    def get_id(self) -> str:
        """Generate unique chunk ID."""
        if not self.id:
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            self.id = f"{self.document_id}_chunk_{self.chunk_index}_{content_hash}"
        return self.id


class ChunkingConfig(BaseModel):
    """Configuration for document chunking."""
    strategy: ChunkingStrategy = Field(default=ChunkingStrategy.RECURSIVE)
    chunk_size: int = Field(default=1000, description="Target chunk size in characters")
    chunk_overlap: int = Field(default=200, description="Overlap between chunks")
    min_chunk_size: int = Field(default=100, description="Minimum chunk size")
    max_chunk_size: int = Field(default=2000, description="Maximum chunk size")
    preserve_sentences: bool = Field(default=True, description="Try to preserve sentence boundaries")
    preserve_paragraphs: bool = Field(default=True, description="Try to preserve paragraph boundaries")


class DocumentProcessor:
    """Base class for document format processors."""
    
    def __init__(self, logger=None):
        self.logger = logger or get_logger(__name__)
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text content from document."""
        raise NotImplementedError
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document metadata."""
        return {}
    
    async def extract_multimodal_content(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract multimodal content (text, images, tables) with layout information."""
        return []


class TextProcessor(DocumentProcessor):
    """Processor for plain text files."""
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin1', 'cp1252', 'ascii']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise AkashaError(f"Could not decode text file: {file_path}")


class PDFProcessor(DocumentProcessor):
    """Processor for PDF files."""
    
    def __init__(self, logger=None):
        super().__init__(logger)
        self._pypdf2 = None
    
    async def _import_pypdf2(self):
        """Lazy import PyPDF2."""
        if self._pypdf2 is None:
            try:
                import PyPDF2
                self._pypdf2 = PyPDF2
            except ImportError:
                raise AkashaError("PyPDF2 not installed. Install with: pip install PyPDF2")
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        await self._import_pypdf2()
        
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = self._pypdf2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
        
        return "\n\n".join(text_content)
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF metadata."""
        await self._import_pypdf2()
        
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = self._pypdf2.PdfReader(file)
                
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_meta.get('/Title', ''),
                        'author': pdf_meta.get('/Author', ''),
                        'subject': pdf_meta.get('/Subject', ''),
                        'creator': pdf_meta.get('/Creator', ''),
                        'producer': pdf_meta.get('/Producer', ''),
                        'creation_date': str(pdf_meta.get('/CreationDate', '')),
                        'modification_date': str(pdf_meta.get('/ModDate', ''))
                    })
                
                metadata['page_count'] = len(pdf_reader.pages)
        except Exception as e:
            self.logger.warning(f"Failed to extract PDF metadata: {e}")
        
        return metadata


class DOCXProcessor(DocumentProcessor):
    """Processor for DOCX files."""
    
    def __init__(self, logger=None):
        super().__init__(logger)
        self._docx = None
    
    async def _import_docx(self):
        """Lazy import python-docx."""
        if self._docx is None:
            try:
                import docx
                self._docx = docx
            except ImportError:
                raise AkashaError("python-docx not installed. Install with: pip install python-docx")
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        await self._import_docx()
        
        doc = self._docx.Document(file_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        return "\n".join(paragraphs)
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        await self._import_docx()
        
        metadata = {}
        try:
            doc = self._docx.Document(file_path)
            core_props = doc.core_properties
            
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'category': core_props.category or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision if core_props.revision else 0,
            })
            
            metadata['paragraph_count'] = len(doc.paragraphs)
        except Exception as e:
            self.logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        return metadata


class MinerU2Processor(DocumentProcessor):
    """Advanced PDF processor using MinerU 2 with OCR fallback."""
    
    def __init__(self, logger=None, enable_ocr=True, ocr_backend="paddleocr"):
        super().__init__(logger)
        self.enable_ocr = enable_ocr
        self.ocr_backend = ocr_backend
        self._mineru = None
        self._ocr_engine = None
        self._image_processors = None
    
    async def _import_mineru(self):
        """Lazy import MinerU."""
        if self._mineru is None:
            try:
                # Import MinerU components
                from magic_pdf.pipe.UNIPipe import UNIPipe
                from magic_pdf.pipe.OCRPipe import OCRPipe
                from magic_pdf.pipe.TXTPipe import TXTParser
                from magic_pdf.model.doc_analyze_by_custom_model import CustomPEKModel
                self._mineru = {
                    'UNIPipe': UNIPipe,
                    'OCRPipe': OCRPipe,
                    'TXTParser': TXTParser,
                    'CustomPEKModel': CustomPEKModel
                }
                self.logger.info("MinerU 2 imported successfully")
            except ImportError as e:
                self.logger.warning(f"MinerU 2 not available: {e}")
                self._mineru = None
    
    async def _import_ocr_engine(self):
        """Lazy import OCR engine."""
        if self._ocr_engine is None and self.enable_ocr:
            try:
                if self.ocr_backend == "paddleocr":
                    from paddleocr import PaddleOCR
                    self._ocr_engine = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
                elif self.ocr_backend == "easyocr":
                    import easyocr
                    self._ocr_engine = easyocr.Reader(['en'])
                else:
                    raise AkashaError(f"Unsupported OCR backend: {self.ocr_backend}")
                
                self.logger.info(f"OCR engine ({self.ocr_backend}) initialized")
            except ImportError as e:
                self.logger.warning(f"OCR backend {self.ocr_backend} not available: {e}")
                self._ocr_engine = None
    
    async def _import_image_processors(self):
        """Lazy import image processing libraries."""
        if self._image_processors is None:
            try:
                from PIL import Image
                import cv2
                self._image_processors = {
                    'PIL': Image,
                    'cv2': cv2
                }
            except ImportError as e:
                self.logger.warning(f"Image processing libraries not available: {e}")
                self._image_processors = None
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text using MinerU 2 with OCR fallback."""
        await self._import_mineru()
        
        # First try MinerU 2 if available
        if self._mineru:
            try:
                return await self._extract_with_mineru(file_path)
            except Exception as e:
                self.logger.warning(f"MinerU 2 extraction failed: {e}, falling back to OCR")
        
        # Fallback to OCR if enabled
        if self.enable_ocr:
            return await self._extract_with_ocr(file_path)
        
        # Final fallback to PyPDF2
        return await self._extract_with_pypdf2(file_path)
    
    async def _extract_with_mineru(self, file_path: Path) -> str:
        """Extract text using MinerU 2."""
        try:
            # Initialize MinerU pipeline
            pipe = self._mineru['UNIPipe'](
                pdf_path=str(file_path),
                output_dir=tempfile.mkdtemp(),
                work_dir=tempfile.mkdtemp()
            )
            
            # Process the PDF
            result = await asyncio.get_event_loop().run_in_executor(
                None, pipe.pipe_analyze
            )
            
            # Extract text content
            text_content = []
            if result and 'pages' in result:
                for page_num, page_data in enumerate(result['pages'], 1):
                    page_text = []
                    if 'blocks' in page_data:
                        for block in page_data['blocks']:
                            if block.get('type') == 'text' and 'text' in block:
                                page_text.append(block['text'])
                    
                    if page_text:
                        text_content.append(f"[Page {page_num}]\n" + "\n".join(page_text))
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            self.logger.error(f"MinerU 2 extraction failed: {e}")
            raise
    
    async def _extract_with_ocr(self, file_path: Path) -> str:
        """Extract text using OCR."""
        await self._import_ocr_engine()
        await self._import_image_processors()
        
        if not self._ocr_engine or not self._image_processors:
            raise AkashaError("OCR engine or image processors not available")
        
        try:
            # Convert PDF pages to images first
            images = await self._pdf_to_images(file_path)
            
            text_content = []
            for page_num, image in enumerate(images, 1):
                # Run OCR on the image
                if self.ocr_backend == "paddleocr":
                    result = self._ocr_engine.ocr(np.array(image), cls=True)
                    page_text = []
                    if result and result[0]:
                        for line in result[0]:
                            if line[1][1] > 0.5:  # Confidence threshold
                                page_text.append(line[1][0])
                elif self.ocr_backend == "easyocr":
                    result = self._ocr_engine.readtext(np.array(image))
                    page_text = []
                    for (bbox, text, confidence) in result:
                        if confidence > 0.5:  # Confidence threshold
                            page_text.append(text)
                
                if page_text:
                    text_content.append(f"[Page {page_num}]\n" + "\n".join(page_text))
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            self.logger.error(f"OCR extraction failed: {e}")
            raise
    
    async def _pdf_to_images(self, file_path: Path) -> List:
        """Convert PDF pages to images."""
        try:
            import fitz  # PyMuPDF
            images = []
            
            def convert_pages():
                doc = fitz.open(str(file_path))
                page_images = []
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x resolution
                    img_data = pix.tobytes("png")
                    image = self._image_processors['PIL'].open(BytesIO(img_data))
                    page_images.append(image)
                doc.close()
                return page_images
            
            images = await asyncio.get_event_loop().run_in_executor(None, convert_pages)
            return images
        
        except ImportError:
            # Fallback to pdf2image if PyMuPDF not available
            try:
                from pdf2image import convert_from_path
                images = await asyncio.get_event_loop().run_in_executor(
                    None, convert_from_path, str(file_path)
                )
                return images
            except ImportError:
                raise AkashaError("Neither PyMuPDF nor pdf2image available for PDF to image conversion")
    
    async def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Fallback extraction using PyPDF2."""
        try:
            import PyPDF2
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            return "\n\n".join(text_content)
        except ImportError:
            raise AkashaError("PyPDF2 not available for fallback extraction")
    
    async def extract_multimodal_content(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract multimodal content with layout information."""
        await self._import_mineru()
        
        if not self._mineru:
            self.logger.warning("MinerU 2 not available, falling back to basic text extraction")
            text = await self.extract_text(file_path)
            return [{"type": "text", "content": text, "page": 1}]
        
        try:
            # Initialize MinerU pipeline for multimodal extraction
            pipe = self._mineru['UNIPipe'](
                pdf_path=str(file_path),
                output_dir=tempfile.mkdtemp(),
                work_dir=tempfile.mkdtemp()
            )
            
            # Process the PDF with full layout analysis
            result = await asyncio.get_event_loop().run_in_executor(
                None, pipe.pipe_analyze
            )
            
            multimodal_content = []
            if result and 'pages' in result:
                for page_num, page_data in enumerate(result['pages'], 1):
                    if 'blocks' in page_data:
                        for block_idx, block in enumerate(page_data['blocks']):
                            content_item = {
                                "page": page_num,
                                "block_index": block_idx,
                                "bbox": block.get('bbox', []),
                                "confidence": block.get('confidence', 1.0)
                            }
                            
                            if block.get('type') == 'text':
                                content_item.update({
                                    "type": "text",
                                    "content": block.get('text', ''),
                                    "content_type": self._classify_text_content(block.get('text', ''))
                                })
                            elif block.get('type') == 'image':
                                content_item.update({
                                    "type": "image",
                                    "content": block.get('image_path', ''),
                                    "image_data": block.get('image_data', ''),
                                    "caption": block.get('caption', '')
                                })
                            elif block.get('type') == 'table':
                                content_item.update({
                                    "type": "table",
                                    "content": block.get('table_html', ''),
                                    "table_data": block.get('table_data', []),
                                    "caption": block.get('caption', '')
                                })
                            elif block.get('type') == 'formula':
                                content_item.update({
                                    "type": "formula",
                                    "content": block.get('latex', ''),
                                    "formula_image": block.get('formula_image', '')
                                })
                            
                            multimodal_content.append(content_item)
            
            return multimodal_content
        
        except Exception as e:
            self.logger.error(f"Multimodal extraction failed: {e}")
            # Fallback to basic text extraction
            text = await self.extract_text(file_path)
            return [{"type": "text", "content": text, "page": 1}]
    
    def _classify_text_content(self, text: str) -> ContentType:
        """Classify text content type based on patterns."""
        text_clean = text.strip()
        
        # Title detection (short, often all caps or title case)
        if len(text_clean) < 100 and (text_clean.isupper() or text_clean.istitle()):
            return ContentType.TITLE
        
        # Header detection (starts with numbers, short lines)
        if len(text_clean) < 200:
            if text_clean.startswith(('1.', '2.', '3.', '4.', '5.', 'Chapter', 'Section')):
                return ContentType.HEADER
            if text_clean.endswith(':'):
                return ContentType.HEADER
        
        # List detection
        if any(text_clean.startswith(marker) for marker in ['•', '-', '*', '1.', '2.', '3.']):
            return ContentType.LIST
        
        # Footnote detection
        if text_clean.startswith(('*', '†', '‡', '§')) or text_clean.startswith(tuple('0123456789')):
            if len(text_clean) < 300:
                return ContentType.FOOTNOTE
        
        # Default to paragraph
        return ContentType.PARAGRAPH
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract comprehensive PDF metadata."""
        metadata = {}
        
        # Try MinerU 2 first for advanced metadata
        await self._import_mineru()
        if self._mineru:
            try:
                pipe = self._mineru['UNIPipe'](
                    pdf_path=str(file_path),
                    output_dir=tempfile.mkdtemp(),
                    work_dir=tempfile.mkdtemp()
                )
                
                result = await asyncio.get_event_loop().run_in_executor(
                    None, pipe.pipe_analyze
                )
                
                if result:
                    metadata.update({
                        'total_pages': len(result.get('pages', [])),
                        'processing_method': 'mineru2',
                        'layout_detected': True,
                        'multimodal_content': True
                    })
                    
                    # Count content types
                    content_stats = {'text': 0, 'image': 0, 'table': 0, 'formula': 0}
                    for page_data in result.get('pages', []):
                        for block in page_data.get('blocks', []):
                            block_type = block.get('type', 'text')
                            if block_type in content_stats:
                                content_stats[block_type] += 1
                    
                    metadata['content_statistics'] = content_stats
                    
            except Exception as e:
                self.logger.warning(f"MinerU 2 metadata extraction failed: {e}")
        
        # Fallback to basic PDF metadata
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_meta.get('/Title', ''),
                        'author': pdf_meta.get('/Author', ''),
                        'subject': pdf_meta.get('/Subject', ''),
                        'creator': pdf_meta.get('/Creator', ''),
                        'producer': pdf_meta.get('/Producer', ''),
                        'creation_date': str(pdf_meta.get('/CreationDate', '')),
                        'modification_date': str(pdf_meta.get('/ModDate', ''))
                    })
                
                if 'total_pages' not in metadata:
                    metadata['total_pages'] = len(pdf_reader.pages)
        
        except Exception as e:
            self.logger.warning(f"Basic PDF metadata extraction failed: {e}")
        
        return metadata


class DocumentChunker:
    """Handles document chunking with various strategies."""
    
    def __init__(self, config: ChunkingConfig, logger=None):
        self.config = config
        self.logger = logger or get_logger(__name__)
    
    async def chunk_text(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Chunk text based on configured strategy."""
        if self.config.strategy == ChunkingStrategy.FIXED_SIZE:
            return await self._chunk_fixed_size(text, document_id)
        elif self.config.strategy == ChunkingStrategy.SENTENCE:
            return await self._chunk_by_sentence(text, document_id)
        elif self.config.strategy == ChunkingStrategy.PARAGRAPH:
            return await self._chunk_by_paragraph(text, document_id)
        elif self.config.strategy == ChunkingStrategy.RECURSIVE:
            return await self._chunk_recursive(text, document_id)
        elif self.config.strategy == ChunkingStrategy.LAYOUT_AWARE:
            # Layout-aware chunking requires multimodal content
            raise AkashaError("Layout-aware chunking requires multimodal content input")
        else:
            raise AkashaError(f"Unsupported chunking strategy: {self.config.strategy}")
    
    async def chunk_multimodal_content(self, multimodal_content: List[Dict[str, Any]], document_id: str) -> List[DocumentChunk]:
        """Chunk multimodal content based on layout information."""
        if self.config.strategy == ChunkingStrategy.LAYOUT_AWARE:
            return await self._chunk_layout_aware(multimodal_content, document_id)
        else:
            # Fallback to text-only chunking
            text_content = []
            for item in multimodal_content:
                if item.get('type') == 'text':
                    text_content.append(item.get('content', ''))
            
            full_text = '\n\n'.join(text_content)
            return await self.chunk_text(full_text, document_id)
    
    async def _chunk_fixed_size(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Chunk text into fixed-size pieces with overlap."""
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + self.config.chunk_size, len(text))
            
            # Try to break at word boundary if preserving sentences
            if self.config.preserve_sentences and end < len(text):
                sentence_end = text.rfind('. ', start, end)
                if sentence_end > start:
                    end = sentence_end + 2
            
            chunk_text = text[start:end].strip()
            if len(chunk_text) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    id="",
                    content=chunk_text,
                    document_id=document_id,
                    chunk_index=chunk_index,
                    start_offset=start,
                    end_offset=end
                )
                chunk.id = chunk.get_id()
                chunks.append(chunk)
                chunk_index += 1
            
            # Move start position considering overlap
            start = max(start + self.config.chunk_size - self.config.chunk_overlap, end)
        
        return chunks
    
    async def _chunk_by_sentence(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Chunk text by sentences, grouping to target size."""
        # Simple sentence splitting - could be enhanced with nltk/spacy
        sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        start_offset = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > self.config.chunk_size and current_chunk:
                # Create chunk from current sentences
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text) >= self.config.min_chunk_size:
                    chunk = DocumentChunk(
                        id="",
                        content=chunk_text,
                        document_id=document_id,
                        chunk_index=chunk_index,
                        start_offset=start_offset,
                        end_offset=start_offset + len(chunk_text)
                    )
                    chunk.id = chunk.get_id()
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_sentences = current_chunk[-self.config.chunk_overlap//100:] if self.config.chunk_overlap > 0 else []
                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s) for s in current_chunk)
                start_offset += len(chunk_text) - sum(len(s) for s in overlap_sentences)
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    id="",
                    content=chunk_text,
                    document_id=document_id,
                    chunk_index=chunk_index,
                    start_offset=start_offset,
                    end_offset=start_offset + len(chunk_text)
                )
                chunk.id = chunk.get_id()
                chunks.append(chunk)
        
        return chunks
    
    async def _chunk_by_paragraph(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Chunk text by paragraphs, grouping to target size."""
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        start_offset = 0
        
        for paragraph in paragraphs:
            paragraph_size = len(paragraph)
            
            if current_size + paragraph_size > self.config.chunk_size and current_chunk:
                # Create chunk from current paragraphs
                chunk_text = '\n\n'.join(current_chunk)
                if len(chunk_text) >= self.config.min_chunk_size:
                    chunk = DocumentChunk(
                        id="",
                        content=chunk_text,
                        document_id=document_id,
                        chunk_index=chunk_index,
                        start_offset=start_offset,
                        end_offset=start_offset + len(chunk_text)
                    )
                    chunk.id = chunk.get_id()
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Start new chunk
                current_chunk = [paragraph]
                current_size = paragraph_size
                start_offset += len(chunk_text)
            else:
                current_chunk.append(paragraph)
                current_size += paragraph_size + 2  # Add 2 for \n\n
        
        # Add final chunk
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            if len(chunk_text) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    id="",
                    content=chunk_text,
                    document_id=document_id,
                    chunk_index=chunk_index,
                    start_offset=start_offset,
                    end_offset=start_offset + len(chunk_text)
                )
                chunk.id = chunk.get_id()
                chunks.append(chunk)
        
        return chunks
    
    async def _chunk_recursive(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Recursive chunking with hierarchical separators."""
        separators = ['\n\n', '\n', '. ', ' ']
        
        def _split_text_recursive(text: str, separators: List[str]) -> List[str]:
            """Recursively split text using hierarchical separators."""
            if not separators or len(text) <= self.config.chunk_size:
                return [text]
            
            separator = separators[0]
            remaining_separators = separators[1:]
            
            splits = text.split(separator)
            good_splits = []
            
            for split in splits:
                if len(split) <= self.config.chunk_size:
                    good_splits.append(split)
                else:
                    # Recursively split large chunks
                    sub_splits = _split_text_recursive(split, remaining_separators)
                    good_splits.extend(sub_splits)
            
            return good_splits
        
        # Get initial splits
        splits = _split_text_recursive(text, separators)
        
        # Combine splits to target size with CHARACTER-BASED overlap
        chunks = []
        current_chunk = []
        chunk_index = 0
        text_position = 0  # Track actual position in original text
        
        i = 0
        while i < len(splits):
            current_chunk = []
            current_content = ""
            chunk_start = text_position
            
            # Build chunk by adding splits until we reach target size
            while i < len(splits):
                split = splits[i]
                # Account for separator when joining (except for first split in chunk)
                separator_len = 1 if current_chunk else 0  # '\n' separator
                
                if current_content and len(current_content) + separator_len + len(split) > self.config.chunk_size:
                    break
                
                if current_chunk:  # Add separator for non-first splits
                    current_content += '\n'
                current_content += split
                current_chunk.append(split)
                i += 1
            
            # Create chunk if we have content
            if current_content.strip() and len(current_content.strip()) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    id="",
                    content=current_content.strip(),
                    document_id=document_id,
                    chunk_index=chunk_index,
                    start_offset=chunk_start,
                    end_offset=chunk_start + len(current_content.strip())
                )
                chunk.id = chunk.get_id()
                chunks.append(chunk)
                chunk_index += 1
                
                # Calculate overlap for next chunk (character-based)
                if i < len(splits) and self.config.chunk_overlap > 0:
                    # Find how much content to overlap (from end of current chunk)
                    content_to_overlap = current_content.strip()
                    overlap_chars = min(self.config.chunk_overlap, len(content_to_overlap))
                    
                    if overlap_chars > 0:
                        # Find where the overlap starts in the original content
                        overlap_start = len(content_to_overlap) - overlap_chars
                        overlap_content = content_to_overlap[overlap_start:]
                        
                        # Move back in the splits to include overlapping content
                        # This is a simple approach - we'll go back one split for overlap
                        if len(current_chunk) > 1:
                            i -= 1  # Go back one split for overlap
                            text_position = chunk_start + len(content_to_overlap) - len(splits[i])
                        else:
                            text_position = chunk_start + len(content_to_overlap)
                    else:
                        text_position = chunk_start + len(content_to_overlap)
                else:
                    text_position = chunk_start + len(current_content.strip())
            else:
                # Skip empty/too-small content
                if i < len(splits):
                    text_position += len(splits[i])
                    i += 1
        
        return chunks
    
    async def _chunk_layout_aware(self, multimodal_content: List[Dict[str, Any]], document_id: str) -> List[DocumentChunk]:
        """Layout-aware chunking based on document structure and multimodal content."""
        chunks = []
        chunk_index = 0
        
        # Group content by page first
        pages = {}
        for item in multimodal_content:
            page_num = item.get('page', 1)
            if page_num not in pages:
                pages[page_num] = []
            pages[page_num].append(item)
        
        # Process each page
        for page_num in sorted(pages.keys()):
            page_items = pages[page_num]
            
            # Sort items by bounding box (top to bottom, left to right)
            def sort_key(item):
                bbox = item.get('bbox', [0, 0, 0, 0])
                if len(bbox) >= 4:
                    return (bbox[1], bbox[0])  # y, x coordinates
                return (0, 0)
            
            page_items.sort(key=sort_key)
            
            # Group items into logical chunks
            current_chunk_items = []
            current_chunk_size = 0
            
            for item in page_items:
                item_content = item.get('content', '')
                item_size = len(item_content)
                item_type = item.get('type', 'text')
                
                # Different handling for different content types
                if item_type == 'text':
                    content_type = item.get('content_type', ContentType.PARAGRAPH)
                    
                    # Headers and titles start new chunks
                    if content_type in [ContentType.TITLE, ContentType.HEADER]:
                        if current_chunk_items:
                            # Create chunk from current items
                            chunk = await self._create_multimodal_chunk(
                                current_chunk_items, document_id, chunk_index, page_num
                            )
                            if chunk:
                                chunks.append(chunk)
                                chunk_index += 1
                        
                        # Start new chunk with header/title
                        current_chunk_items = [item]
                        current_chunk_size = item_size
                    else:
                        # Check if adding this item would exceed chunk size
                        if (current_chunk_size + item_size > self.config.chunk_size and 
                            current_chunk_items):
                            # Create chunk from current items
                            chunk = await self._create_multimodal_chunk(
                                current_chunk_items, document_id, chunk_index, page_num
                            )
                            if chunk:
                                chunks.append(chunk)
                                chunk_index += 1
                            
                            # Start new chunk with overlap if configured
                            if self.config.chunk_overlap > 0 and current_chunk_items:
                                # Keep last item for overlap
                                overlap_items = current_chunk_items[-1:]
                                overlap_size = sum(len(i.get('content', '')) for i in overlap_items)
                                current_chunk_items = overlap_items + [item]
                                current_chunk_size = overlap_size + item_size
                            else:
                                current_chunk_items = [item]
                                current_chunk_size = item_size
                        else:
                            current_chunk_items.append(item)
                            current_chunk_size += item_size
                
                elif item_type in ['image', 'table', 'formula']:
                    # Non-text items always start new chunks
                    if current_chunk_items:
                        # Create chunk from current text items
                        chunk = await self._create_multimodal_chunk(
                            current_chunk_items, document_id, chunk_index, page_num
                        )
                        if chunk:
                            chunks.append(chunk)
                            chunk_index += 1
                    
                    # Create dedicated chunk for multimodal item
                    chunk = await self._create_multimodal_chunk(
                        [item], document_id, chunk_index, page_num
                    )
                    if chunk:
                        chunks.append(chunk)
                        chunk_index += 1
                    
                    current_chunk_items = []
                    current_chunk_size = 0
            
            # Create final chunk for remaining items
            if current_chunk_items:
                chunk = await self._create_multimodal_chunk(
                    current_chunk_items, document_id, chunk_index, page_num
                )
                if chunk:
                    chunks.append(chunk)
                    chunk_index += 1
        
        return chunks
    
    async def _create_multimodal_chunk(self, items: List[Dict[str, Any]], 
                                     document_id: str, chunk_index: int, 
                                     page_num: int) -> Optional[DocumentChunk]:
        """Create a DocumentChunk from multimodal content items."""
        if not items:
            return None
        
        # Combine text content
        text_parts = []
        chunk_metadata = {
            'page_number': page_num,
            'item_count': len(items),
            'content_types': []
        }
        
        # Handle different content types
        for item in items:
            item_type = item.get('type', 'text')
            content = item.get('content', '')
            
            if item_type == 'text':
                if content.strip():
                    text_parts.append(content)
                    chunk_metadata['content_types'].append(item.get('content_type', 'text'))
            
            elif item_type == 'image':
                # Add image placeholder and caption
                caption = item.get('caption', '')
                image_desc = f"[IMAGE: {caption}]" if caption else "[IMAGE]"
                text_parts.append(image_desc)
                chunk_metadata['content_types'].append('image')
                chunk_metadata['has_image'] = True
                
                # Store image data if available
                if item.get('image_data'):
                    chunk_metadata['image_data'] = item.get('image_data')
            
            elif item_type == 'table':
                # Add table content
                table_content = item.get('content', '')
                caption = item.get('caption', '')
                if caption:
                    text_parts.append(f"[TABLE: {caption}]")
                if table_content:
                    text_parts.append(table_content)
                chunk_metadata['content_types'].append('table')
                chunk_metadata['has_table'] = True
            
            elif item_type == 'formula':
                # Add formula content
                formula_content = item.get('content', '')
                if formula_content:
                    text_parts.append(f"[FORMULA: {formula_content}]")
                chunk_metadata['content_types'].append('formula')
                chunk_metadata['has_formula'] = True
        
        # Create chunk content
        chunk_content = '\n\n'.join(text_parts).strip()
        
        if not chunk_content or len(chunk_content) < self.config.min_chunk_size:
            return None
        
        # Determine primary content type
        primary_content_type = ContentType.TEXT
        if len(items) == 1:
            item_type = items[0].get('type', 'text')
            if item_type == 'image':
                primary_content_type = ContentType.IMAGE
            elif item_type == 'table':
                primary_content_type = ContentType.TABLE
            elif item_type == 'formula':
                primary_content_type = ContentType.FORMULA
            elif item_type == 'text':
                primary_content_type = items[0].get('content_type', ContentType.PARAGRAPH)
        
        # Calculate bounding boxes
        bboxes = [item.get('bbox', []) for item in items if item.get('bbox')]
        combined_bbox = None
        if bboxes:
            # Combine bounding boxes (min x, min y, max x, max y)
            min_x = min(bbox[0] for bbox in bboxes if len(bbox) >= 4)
            min_y = min(bbox[1] for bbox in bboxes if len(bbox) >= 4)
            max_x = max(bbox[2] for bbox in bboxes if len(bbox) >= 4)
            max_y = max(bbox[3] for bbox in bboxes if len(bbox) >= 4)
            combined_bbox = [min_x, min_y, max_x, max_y]
        
        # Calculate confidence
        confidences = [item.get('confidence', 1.0) for item in items]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 1.0
        
        # Create chunk
        chunk = DocumentChunk(
            id="",
            content=chunk_content,
            document_id=document_id,
            chunk_index=chunk_index,
            content_type=primary_content_type,
            page_number=page_num,
            bbox=combined_bbox,
            confidence=avg_confidence,
            metadata=chunk_metadata
        )
        
        chunk.id = chunk.get_id()
        return chunk


class DocumentIngestion:
    """Main document ingestion pipeline."""
    
    def __init__(self, chunking_config: ChunkingConfig = None, logger=None):
        self.logger = logger or get_logger(__name__)
        self.chunking_config = chunking_config or ChunkingConfig()
        self.chunker = DocumentChunker(self.chunking_config, self.logger)
        
        # Initialize processors
        self.processors = {
            DocumentFormat.TEXT: TextProcessor(self.logger),
            DocumentFormat.PDF: MinerU2Processor(self.logger, enable_ocr=True),
            DocumentFormat.DOCX: DOCXProcessor(self.logger),
            DocumentFormat.MARKDOWN: TextProcessor(self.logger),  # Treat as text for now
            DocumentFormat.HTML: TextProcessor(self.logger),      # Could add HTML parser later
        }
        
        # Keep legacy PDF processor as fallback
        self.legacy_pdf_processor = PDFProcessor(self.logger)
    
    def _detect_format(self, file_path: Path) -> DocumentFormat:
        """Detect document format from file extension and MIME type."""
        suffix = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Map extensions to formats
        extension_map = {
            '.txt': DocumentFormat.TEXT,
            '.pdf': DocumentFormat.PDF,
            '.docx': DocumentFormat.DOCX,
            '.doc': DocumentFormat.DOCX,  # Treat .doc as .docx for now
            '.md': DocumentFormat.MARKDOWN,
            '.markdown': DocumentFormat.MARKDOWN,
            '.html': DocumentFormat.HTML,
            '.htm': DocumentFormat.HTML,
            '.json': DocumentFormat.JSON,
            '.csv': DocumentFormat.CSV,
        }
        
        if suffix in extension_map:
            return extension_map[suffix]
        
        # Fallback to MIME type detection
        if mime_type:
            if mime_type.startswith('text/'):
                return DocumentFormat.TEXT
            elif mime_type == 'application/pdf':
                return DocumentFormat.PDF
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return DocumentFormat.DOCX
        
        # Default to text
        return DocumentFormat.TEXT
    
    def _generate_document_id(self, file_path: Path) -> str:
        """Generate unique document ID."""
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()
        return f"doc_{file_hash[:16]}"
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate file content hash for change detection."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    async def process_file(self, file_path: Union[str, Path]) -> tuple[DocumentMetadata, List[DocumentChunk]]:
        """Process a single file and return metadata and chunks."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise AkashaError(f"File not found: {file_path}")
        
        start_time = time.time()
        
        async with PerformanceLogger(f"document_ingestion:{file_path.name}", self.logger):
            # Detect format and get processor
            doc_format = self._detect_format(file_path)
            processor = self.processors.get(doc_format)
            
            if not processor:
                raise AkashaError(f"Unsupported document format: {doc_format}")
            
            # Generate document ID and calculate hash
            document_id = self._generate_document_id(file_path)
            file_hash = self._calculate_file_hash(file_path)
            
            # Extract text content
            try:
                text_content = await processor.extract_text(file_path)
            except Exception as e:
                raise AkashaError(f"Failed to extract text from {file_path}: {e}")
            
            if not text_content.strip():
                raise AkashaError(f"No text content extracted from {file_path}")
            
            # Extract additional metadata
            format_metadata = await processor.extract_metadata(file_path)
            
            # Chunk the document
            chunks = await self.chunker.chunk_text(text_content, document_id)
            
            processing_time = time.time() - start_time
            
            # Create document metadata
            file_stat = file_path.stat()
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            metadata = DocumentMetadata(
                document_id=document_id,
                file_path=str(file_path),
                file_name=file_path.name,
                file_size=file_stat.st_size,
                file_hash=file_hash,
                mime_type=mime_type or "application/octet-stream",
                format=doc_format,
                processed_at=time.time(),
                chunk_count=len(chunks),
                processing_time=processing_time,
                custom_metadata=format_metadata
            )
            
            self.logger.info(
                "Document processed successfully",
                document_id=document_id,
                file_name=file_path.name,
                format=doc_format.value,
                chunk_count=len(chunks),
                processing_time=processing_time
            )
            
            return metadata, chunks
    
    async def process_directory(self, directory_path: Union[str, Path], 
                               recursive: bool = True,
                               file_patterns: List[str] = None) -> Iterator[tuple[DocumentMetadata, List[DocumentChunk]]]:
        """Process all supported files in a directory."""
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise AkashaError(f"Directory not found: {directory_path}")
        
        # Default file patterns
        if file_patterns is None:
            file_patterns = ["*.txt", "*.pdf", "*.docx", "*.md", "*.html"]
        
        # Collect files
        files_to_process = []
        for pattern in file_patterns:
            if recursive:
                files_to_process.extend(directory_path.rglob(pattern))
            else:
                files_to_process.extend(directory_path.glob(pattern))
        
        self.logger.info(
            "Starting directory processing",
            directory=str(directory_path),
            file_count=len(files_to_process),
            recursive=recursive
        )
        
        # Process files
        for file_path in files_to_process:
            try:
                yield await self.process_file(file_path)
            except Exception as e:
                self.logger.error(
                    "Failed to process file",
                    file_path=str(file_path),
                    error=str(e)
                )
                # Continue with other files
                continue
    
    async def batch_process_files(self, file_paths: List[Union[str, Path]], 
                                  max_concurrent: int = 5) -> List[tuple[DocumentMetadata, List[DocumentChunk]]]:
        """Process multiple files concurrently."""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_with_semaphore(file_path):
            async with semaphore:
                return await self.process_file(file_path)
        
        tasks = [process_with_semaphore(fp) for fp in file_paths]
        results = []
        
        for task in asyncio.as_completed(tasks):
            try:
                result = await task
                results.append(result)
            except Exception as e:
                self.logger.error(f"Failed to process file in batch: {e}")
                # Continue with other files
                continue
        
        return results